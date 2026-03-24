import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

def extract_text(docx_path, txt_path):
    doc = docx.Document(docx_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
            
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text.append(" | ".join(row_text))

    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(text))

if __name__ == '__main__':
    docx_file = r"c:\Users\reddy\OneDrive\Desktop\Portfolio\Nethrananda-UpdatedCV.doc (1).docx"
    txt_file = r"c:\Users\reddy\OneDrive\Desktop\Portfolio\Nethrananda-UpdatedCV.txt"
    extract_text(docx_file, txt_file)
    print(f"Extracted to {txt_file}")
